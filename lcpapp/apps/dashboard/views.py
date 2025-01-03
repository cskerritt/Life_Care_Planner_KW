from datetime import datetime, timedelta
import os
from docx import Document
from docx.shared import Inches

from django.contrib.auth.decorators import login_required
from django.template.response import TemplateResponse
from django.utils import timezone
from django.shortcuts import get_object_or_404, redirect, render
from django.http import Http404, HttpResponse, JsonResponse
from django.urls import reverse
from django.conf import settings
from django.views.decorators.csrf import csrf_exempt
from django.utils.text import slugify
from django.db import models
from datetime import datetime


from apps.api.models import (
    Evaluee, Task, CarePlan, CarePlanItem
)


@login_required
def create_care_plan(request):
    """
    Create a new life care plan
    """
    # Check for evaluee_id in GET parameters (from evaluee list)
    initial_evaluee_id = request.GET.get('evaluee_id')
    initial_evaluee = None
    if initial_evaluee_id:
        initial_evaluee = get_object_or_404(Evaluee, id=initial_evaluee_id, user=request.user)

    if request.method == "POST":
        evaluee_id = request.POST.get('evaluee_id')
        title = request.POST.get('title', 'New Life Care Plan')
        
        evaluee = get_object_or_404(Evaluee, id=evaluee_id, user=request.user)
        care_plan = CarePlan.objects.create(
            evaluee=evaluee,
            title=title,
            description="",
            start_date=timezone.now().date(),
            status='draft'
        )
        return redirect('dashboard:care_plan_edit', plan_id=care_plan.id)
    
    # Get list of evaluees for selection
    evaluees = Evaluee.objects.filter(user=request.user)
    return TemplateResponse(
        request,
        "dashboard/care_plan_create.html",
        context={
            "active_tab": "life-care-plans",
            "evaluees": evaluees,
            "initial_evaluee": initial_evaluee,
        },
    )

@login_required
def evaluee_list(request):
    """
    Display list of all evaluees
    """
    evaluees = Evaluee.objects.filter(user=request.user).order_by('-created_at')
    return TemplateResponse(
        request,
        "dashboard/evaluee_list.html",
        context={
            "active_tab": "evaluees",
            "evaluees": evaluees,
        },
    )

@login_required
def evaluee_create(request):
    """
    Create a new evaluee
    """
    if request.method == "POST":
        evaluee = Evaluee.objects.create(
            user=request.user,
            first_name=request.POST.get('first_name'),
            last_name=request.POST.get('last_name'),
            date_of_birth=request.POST.get('date_of_birth'),
            medical_record_number=request.POST.get('medical_record_number'),
            primary_diagnosis=request.POST.get('primary_diagnosis'),
            notes=request.POST.get('notes', '')
        )
        return redirect('dashboard:evaluee_list')
    
    return TemplateResponse(
        request,
        "dashboard/evaluee_create.html",
        context={
            "active_tab": "evaluees",
        },
    )

@login_required
def evaluee_detail(request, evaluee_id):
    """
    Display detailed view of a specific evaluee
    """
    evaluee = get_object_or_404(Evaluee, id=evaluee_id, user=request.user)
    care_plans = CarePlan.objects.filter(evaluee=evaluee).order_by('-created_at')
    
    return TemplateResponse(
        request,
        "dashboard/evaluee_detail.html",
        context={
            "active_tab": "evaluees",
            "evaluee": evaluee,
            "care_plans": care_plans,
        },
    )

@login_required
def evaluee_edit(request, evaluee_id):
    """
    Edit a specific evaluee
    """
    evaluee = get_object_or_404(Evaluee, id=evaluee_id, user=request.user)
    
    if request.method == "POST":
        evaluee.first_name = request.POST.get('first_name')
        evaluee.last_name = request.POST.get('last_name')
        evaluee.date_of_birth = request.POST.get('date_of_birth')
        evaluee.medical_record_number = request.POST.get('medical_record_number')
        evaluee.primary_diagnosis = request.POST.get('primary_diagnosis')
        evaluee.notes = request.POST.get('notes', '')
        evaluee.save()
        return redirect('dashboard:evaluee_detail', evaluee_id=evaluee_id)
    
    return TemplateResponse(
        request,
        "dashboard/evaluee_edit.html",
        context={
            "active_tab": "evaluees",
            "evaluee": evaluee,
        },
    )

@login_required
def dashboard(request):
    """
    Main dashboard view showing overview of evaluees and life care plans
    """
    # Get current user's evaluees
    evaluees = Evaluee.objects.filter(user=request.user)
    
    # Get upcoming tasks across all life care plans
    upcoming_tasks = Task.objects.filter(
        care_plan__evaluee__user=request.user,
        status__in=['pending', 'in_progress'],
        due_date__gte=timezone.now()
    ).order_by('due_date')[:5]
    
    # Get active life care plans
    active_care_plans = CarePlan.objects.filter(
        evaluee__user=request.user,
        status='active'
    ).order_by('-updated_at')[:5]
    
    # Get recent items
    recent_items = CarePlanItem.objects.filter(
        care_plan__evaluee__user=request.user
    ).order_by('-created_at')[:5]
    
    # Get tasks due today
    today = timezone.now().date()
    tasks_due_today = Task.objects.filter(
        care_plan__evaluee__user=request.user,
        status__in=['pending', 'in_progress'],
        due_date__date=today
    ).order_by('due_date')
    
    # Get overdue tasks
    overdue_tasks = Task.objects.filter(
        care_plan__evaluee__user=request.user,
        status__in=['pending', 'in_progress'],
        due_date__lt=timezone.now()
    ).order_by('due_date')
    
    # Calculate task statistics
    total_tasks = Task.objects.filter(care_plan__evaluee__user=request.user).count()
    completed_tasks = Task.objects.filter(
        care_plan__evaluee__user=request.user,
        status='completed'
    ).count()
    completion_rate = (completed_tasks / total_tasks * 100) if total_tasks > 0 else 0
    
    return TemplateResponse(
        request,
        "dashboard/user_dashboard.html",
        context={
            "active_tab": "dashboard",
            "evaluees": evaluees,
            "upcoming_tasks": upcoming_tasks,
            "active_care_plans": active_care_plans,
            "recent_items": recent_items,
            "tasks_due_today": tasks_due_today,
            "overdue_tasks": overdue_tasks,
            "total_tasks": total_tasks,
            "completed_tasks": completed_tasks,
            "completion_rate": round(completion_rate, 1),
        },
    )


@login_required
def export_care_plan(request, plan_id):
    """
    Export life care plan details to Word document
    """
    care_plan = get_object_or_404(CarePlan, id=plan_id, evaluee__user=request.user)
    
    # Create Word document
    doc = Document()
    doc.add_heading(f'Life Care Plan: {care_plan.title}', 0)
    
    # Evaluee Information
    doc.add_heading('Evaluee Information', level=1)
    doc.add_paragraph(f'Name: {care_plan.evaluee.first_name} {care_plan.evaluee.last_name}')
    doc.add_paragraph(f'Medical Record Number: {care_plan.evaluee.medical_record_number}')
    doc.add_paragraph(f'Primary Diagnosis: {care_plan.evaluee.primary_diagnosis}')
    
    # Plan Details
    doc.add_heading('Plan Details', level=1)
    doc.add_paragraph(f'Status: {care_plan.get_status_display()}')
    doc.add_paragraph(f'Start Date: {care_plan.start_date}')
    if care_plan.end_date:
        doc.add_paragraph(f'End Date: {care_plan.end_date}')
    
    # Group items by category
    items_by_category = {}
    for item in care_plan.items.all():
        category = item.get_category_display()
        if category not in items_by_category:
            items_by_category[category] = []
        items_by_category[category].append(item)
    
    # Add each category to the document
    for category, items in items_by_category.items():
        doc.add_heading(category, level=1)
        for item in items:
            doc.add_paragraph(f'• {item.title}')
            doc.add_paragraph(f'  Description: {item.description}')
            doc.add_paragraph(f'  Frequency: {item.frequency}')
            doc.add_paragraph(f'  Cost: ${item.cost:,.2f}')
            if item.start_date:
                doc.add_paragraph(f'  Start Date: {item.start_date}')
            if item.end_date:
                doc.add_paragraph(f'  End Date: {item.end_date}')
    
    # Medical Costs
    doc.add_heading('Medical Costs Summary', level=1)
    doc.add_paragraph(f'Total Medical Costs: ${care_plan.total_medical_costs:,.2f}')
    doc.add_paragraph(f'Monthly Cost Average: ${care_plan.monthly_cost_average:,.2f}')
    
    # Life Expectancy
    if care_plan.life_expectancy:
        doc.add_heading('Life Expectancy Estimate', level=1)
        doc.add_paragraph(f'Estimated Years: {care_plan.life_expectancy}')
    
    # Save to temporary file and serve
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename=life_care_plan_{plan_id}.docx'
    doc.save(response)
    
    return response

@login_required
def care_plan_list(request):
    """
    Display list of all life care plans
    """
    care_plans = CarePlan.objects.filter(
        evaluee__user=request.user
    ).select_related('evaluee').order_by('-created_at')
    
    return TemplateResponse(
        request,
        "dashboard/care_plan_list.html",
        context={
            "active_tab": "life-care-plans",
            "care_plans": care_plans,
        },
    )

@login_required
def care_plan_detail(request, plan_id):
    """
    Display detailed view of a specific life care plan
    """
    care_plan = get_object_or_404(CarePlan, id=plan_id, evaluee__user=request.user)
    
    # Group items by category
    items_by_category = {}
    for item in care_plan.items.all():
        category = item.get_category_display()
        if category not in items_by_category:
            items_by_category[category] = []
        items_by_category[category].append(item)
    
    return TemplateResponse(
        request,
        "dashboard/care_plan_detail.html",
        context={
            "active_tab": "life-care-plans",
            "care_plan": care_plan,
            "items_by_category": items_by_category,
        },
    )

@login_required
def care_plan_edit(request, plan_id):
    """
    Edit a specific life care plan
    """
    care_plan = get_object_or_404(CarePlan, id=plan_id, evaluee__user=request.user)
    
    if request.method == "POST":
        # Update basic life care plan info
        care_plan.title = request.POST.get('title')
        care_plan.status = request.POST.get('status')
        care_plan.start_date = request.POST.get('start_date')
        care_plan.end_date = request.POST.get('end_date') or None
        
        # Handle items
        care_plan.items.all().delete()  # Remove existing items
        
        # Get all the arrays from POST
        item_titles = request.POST.getlist('item_title[]')
        item_categories = request.POST.getlist('item_category[]')
        item_descriptions = request.POST.getlist('item_description[]')
        item_frequencies = request.POST.getlist('item_frequency[]')
        item_costs = request.POST.getlist('item_cost[]')
        item_start_dates = request.POST.getlist('item_start_date[]')
        item_end_dates = request.POST.getlist('item_end_date[]')
        item_statuses = request.POST.getlist('item_status[]')
        
        # Create new items
        total_costs = 0
        for i in range(len(item_titles)):
            if item_titles[i]:  # Only create if title exists
                cost = float(item_costs[i] or 0)
                total_costs += cost
                
                CarePlanItem.objects.create(
                    care_plan=care_plan,
                    title=item_titles[i],
                    category=item_categories[i],
                    description=item_descriptions[i],
                    frequency=item_frequencies[i],
                    cost=cost,
                    start_date=item_start_dates[i],
                    end_date=item_end_dates[i] or None,
                    status=item_statuses[i]
                )
        
        # Update care plan costs
        care_plan.total_medical_costs = total_costs
        care_plan.monthly_cost_average = total_costs / 12  # Simple average, could be more sophisticated
        
        care_plan.save()
        return redirect('dashboard:care_plan_detail', plan_id=plan_id)
    
    # Group items by category for display
    items_by_category = {}
    for item in care_plan.items.all():
        category = item.get_category_display()
        if category not in items_by_category:
            items_by_category[category] = []
        items_by_category[category].append(item)
    
    return TemplateResponse(
        request,
        "dashboard/care_plan_edit.html",
        context={
            "active_tab": "life-care-plans",
            "care_plan": care_plan,
            "items_by_category": items_by_category,
            "categories": CarePlanItem.CATEGORY_CHOICES,
        },
    )

@login_required
def medical_costs(request, plan_id=None):
    """
    Medical costs calculator view
    """
    care_plan = None
    if plan_id:
        care_plan = get_object_or_404(CarePlan, id=plan_id, evaluee__user=request.user)
    
    if request.method == "POST" and care_plan:
        # Calculate total costs from all items
        total_costs = care_plan.items.aggregate(models.Sum('cost'))['cost__sum'] or 0
        monthly_costs = total_costs / 12  # Simple average, could be more sophisticated
        
        # Update care plan
        care_plan.total_medical_costs = total_costs
        care_plan.monthly_cost_average = monthly_costs
        care_plan.save()
        
        return redirect('dashboard:care_plan_detail', plan_id=plan_id)
    
    return TemplateResponse(
        request,
        "dashboard/medical_costs.html",
        context={
            "active_tab": "medical-costs",
            "care_plan": care_plan,
        },
    )

@login_required
def life_expectancy(request, plan_id=None):
    """
    Life expectancy calculator view
    """
    care_plan = None
    if plan_id:
        care_plan = get_object_or_404(CarePlan, id=plan_id, evaluee__user=request.user)
    
    if request.method == "POST" and care_plan:
        # Calculate life expectancy
        age = int(request.POST.get('age', 0))
        gender = request.POST.get('gender')
        smoking = request.POST.get('smoking')
        activity = request.POST.get('activity')
        parents_age = int(request.POST.get('parents_age', 0))
        grandparents_age = int(request.POST.get('grandparents_age', 0))
        
        # Base life expectancy
        base_expectancy = 81 if gender == 'female' else 76
        
        # Health impact
        health_impact = 0
        if request.POST.get('condition_diabetes'): health_impact -= 5
        if request.POST.get('condition_heart'): health_impact -= 7
        if request.POST.get('condition_cancer'): health_impact -= 8
        if request.POST.get('condition_respiratory'): health_impact -= 6
        
        # Lifestyle impact
        lifestyle_impact = 0
        if smoking == 'current': lifestyle_impact -= 10
        elif smoking == 'former': lifestyle_impact -= 5
        
        if activity == 'active': lifestyle_impact += 5
        elif activity == 'sedentary': lifestyle_impact -= 3
        
        # Genetic impact
        genetic_impact = 0
        if parents_age and grandparents_age:
            family_longevity = (parents_age + grandparents_age) / 2
            genetic_impact = 5 if family_longevity > 80 else (2 if family_longevity > 70 else -2)
        
        # Calculate final estimate
        estimate = max(age + 1, base_expectancy + health_impact + lifestyle_impact + genetic_impact)
        
        # Update care plan
        care_plan.life_expectancy = round(estimate)
        care_plan.save()
        
        return redirect('dashboard:care_plan_detail', plan_id=plan_id)
    
    return TemplateResponse(
        request,
        "dashboard/life_expectancy.html",
        context={
            "active_tab": "life-expectancy",
            "care_plan": care_plan,
        },
    )
